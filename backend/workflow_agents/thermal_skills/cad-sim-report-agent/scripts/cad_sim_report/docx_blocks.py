from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate

try:
    from PIL import Image
except Exception:
    Image = None


DocxBlock = dict[str, Any]
EAST_ASIA_FONT = "宋体"
LATIN_FONT = "Times New Roman"
BODY_FONT_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(10.5)
CAPTION_FONT_SIZE = Pt(10.5)


def set_run_font(run: Any, size: Any = BODY_FONT_SIZE, bold: bool | None = None, font_name: str = LATIN_FONT) -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font_name)
    run.font.size = size
    if bold is not None:
        run.bold = bold


def apply_paragraph_format(paragraph_obj: Any, first_line_indent: bool = False, align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph_obj.alignment = align
    paragraph_format = paragraph_obj.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.5
    if first_line_indent:
        paragraph_format.first_line_indent = Pt(24)


def set_table_borders(table_obj: Any) -> None:
    tbl_pr = table_obj._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, size in {
        "top": "8",
        "left": "8",
        "bottom": "8",
        "right": "8",
        "insideH": "4",
        "insideV": "4",
    }.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")


def add_text_runs(paragraph_obj: Any, text: str) -> None:
    pattern = re.compile(r"`([^`]+)`")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            append_plain_runs(paragraph_obj, text[pos:match.start()])
        run = paragraph_obj.add_run(match.group(1))
        set_run_font(run, size=BODY_FONT_SIZE, font_name="Consolas")
        pos = match.end()
    if pos < len(text):
        append_plain_runs(paragraph_obj, text[pos:])


def append_plain_runs(paragraph_obj: Any, text: str) -> None:
    parts = text.split("\n")
    for index, part in enumerate(parts):
        if index:
            paragraph_obj.add_run().add_break()
        if part:
            run = paragraph_obj.add_run(part)
            set_run_font(run)


def image_size(path: Path, max_width_in: float) -> tuple[float, float | None]:
    if Image is None:
        return max_width_in, None
    try:
        with Image.open(path) as image:
            width_px, height_px = image.size
    except Exception:
        return max_width_in, None
    width_in = min(max_width_in, width_px / 140)
    height_in = width_in * height_px / max(width_px, 1)
    return width_in, height_in


def add_image_to_paragraph(paragraph_obj: Any, image_path: Path, max_width_in: float) -> None:
    if not image_path.exists():
        paragraph_obj.add_run(f"[missing image: {image_path.name}]")
        return
    width_in, height_in = image_size(image_path, max_width_in)
    kwargs: dict[str, Any] = {"width": Inches(width_in)}
    if height_in is not None:
        kwargs["height"] = Inches(height_in)
    paragraph_obj.add_run().add_picture(str(image_path), **kwargs)


def add_paragraph_with_inline_code(doc: Any, text: str, style: str | None = None) -> None:
    paragraph_obj = doc.add_paragraph(style=style)
    apply_paragraph_format(paragraph_obj, first_line_indent=style is None)
    add_text_runs(paragraph_obj, text)


def add_caption(doc: Any, text: str) -> None:
    if not text:
        return
    paragraph_obj = doc.add_paragraph()
    apply_paragraph_format(paragraph_obj, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph_obj.add_run(text)
    set_run_font(run, size=CAPTION_FONT_SIZE, bold=False)


def set_cell_width(cell: Any, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def estimate_column_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    column_count = max([len(headers), *(len(row) for row in rows)] or [1])
    weights: list[int] = []
    all_rows = [headers, *rows]
    for index in range(column_count):
        max_len = max((len(str(row[index])) if index < len(row) else 0) for row in all_rows)
        weights.append(max(8, min(max_len, 42)))
    total = sum(weights) or 1
    table_width = 9000
    return [max(900, int(table_width * weight / total)) for weight in weights]


def add_table_from_rows(doc: Any, headers: list[str], rows: list[list[str]], caption: str = "") -> None:
    all_rows = [headers, *rows] if headers else rows
    if not all_rows:
        return
    add_caption(doc, caption)
    column_count = max(len(row) for row in all_rows)
    widths = estimate_column_widths(headers, rows)
    table_obj = doc.add_table(rows=len(all_rows), cols=column_count)
    table_obj.style = "Table Grid"
    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_obj.autofit = False
    set_table_borders(table_obj)
    for r_index, row in enumerate(all_rows):
        for c_index in range(column_count):
            cell = table_obj.cell(r_index, c_index)
            set_cell_width(cell, widths[c_index] if c_index < len(widths) else widths[-1])
            text = row[c_index] if c_index < len(row) else ""
            paragraph_obj = cell.paragraphs[0]
            apply_paragraph_format(paragraph_obj, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            add_text_runs(paragraph_obj, str(text))
            for run in paragraph_obj.runs:
                set_run_font(run, size=TABLE_FONT_SIZE, bold=bool(r_index == 0 and headers))
            if r_index == 0 and headers:
                for run in paragraph_obj.runs:
                    run.bold = True
    doc.add_paragraph("")


def add_image_gallery(doc: Any, images: list[dict[str, Any]], caption: str = "", captions: list[str] | None = None) -> None:
    existing = [image for image in images if image.get("exists")]
    if not existing:
        add_paragraph_with_inline_code(doc, "未找到可用图片。")
        return
    for index, image in enumerate(existing, start=1):
        path = Path(str(image["path"]))
        paragraph_obj = doc.add_paragraph()
        paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_to_paragraph(paragraph_obj, path, 5.8)
        image_caption = captions[index - 1] if captions and index <= len(captions) else caption
        add_caption(doc, image_caption)
    doc.add_paragraph("")


def render_blocks_to_subdoc(template: DocxTemplate, blocks: list[DocxBlock]) -> Any:
    subdoc = template.new_subdoc()
    for block in blocks:
        block_type = block.get("type")
        if block_type == "blank":
            subdoc.add_paragraph("")
        elif block_type == "heading":
            level = min(int(block.get("level", 3)), 3)
            paragraph_obj = subdoc.add_paragraph(str(block.get("text", "")), style=f"Heading {level}")
            apply_paragraph_format(paragraph_obj, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            for run in paragraph_obj.runs:
                set_run_font(run, size=BODY_FONT_SIZE, bold=True)
        elif block_type == "paragraph":
            add_paragraph_with_inline_code(subdoc, str(block.get("text", "")))
        elif block_type == "note":
            add_paragraph_with_inline_code(subdoc, str(block.get("text", "")))
        elif block_type == "list_item":
            style = "List Number" if block.get("ordered") else "List Bullet"
            add_paragraph_with_inline_code(subdoc, str(block.get("text", "")), style=style)
        elif block_type == "table":
            add_table_from_rows(
                subdoc,
                [str(item) for item in block.get("headers", [])],
                [[str(item) for item in row] for row in block.get("rows", [])],
                str(block.get("caption", "")),
            )
        elif block_type == "image_gallery":
            captions = block.get("captions")
            add_image_gallery(
                subdoc,
                block.get("images", []),
                str(block.get("caption", "")),
                captions if isinstance(captions, list) else None,
            )
        else:
            add_paragraph_with_inline_code(subdoc, str(block.get("text", "")))
    return subdoc
